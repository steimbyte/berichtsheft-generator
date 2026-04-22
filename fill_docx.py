#!/usr/bin/env python3
"""
Berichtsheft DOCX Updater
Ersetzt Platzhalter in der aktuellen KW-DOCX und passt automatisch die Größe an.

Usage:
    python3 fill_docx.py --kw 10 --jahr 2026 --nr 130 --abteilung "Betrieb" \
        --montag "Text" --dienstag "Text" --mittwoch "Text" \
        --donnerstag "Text" --freitag "Text" --thema "Text" \
        --datum "02.03.2026 - 08.03.2026"
"""

import argparse
import re
import sys
import zipfile
import subprocess
from pathlib import Path
from docx import Document
from lxml import etree

def replace_in_run(run, replacements):
    """Ersetzt Platzhalter in einem Run (case-insensitive)."""
    if not run.text:
        return
    text = run.text
    for placeholder, value in replacements.items():
        if re.search(re.escape(placeholder), text, re.IGNORECASE):
            text = re.sub(re.escape(placeholder), value, text, flags=re.IGNORECASE)
    run.text = text

def replace_in_paragraph(paragraph, replacements):
    """Ersetzt Platzhalter in allen Runs eines Paragraphs."""
    for run in paragraph.runs:
        replace_in_run(run, replacements)

def replace_in_cell(cell, replacements):
    """Ersetzt Platzhalter in allen Runs einer Tabellenzelle."""
    for para in cell.paragraphs:
        replace_in_paragraph(para, replacements)

def has_placeholders(docx_path):
    """Prüft ob noch Platzhalter im Dokument vorhanden sind."""
    with zipfile.ZipFile(docx_path, 'r') as z:
        xml_content = z.read('word/document.xml')
    root = etree.fromstring(xml_content)
    for t in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
        if t.text and '{' in t.text:
            return True
    return False

def shrink_xml(docx_path, factor=0.95):
    """Verkleinert Zeilenhöhen im XML."""
    with zipfile.ZipFile(docx_path, 'r') as z:
        xml_content = z.read('word/document.xml')
    
    root = etree.fromstring(xml_content)
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    
    modified = 0
    for trHeight in root.iter(f'{{{ns}}}trHeight'):
        val = trHeight.get(f'{{{ns}}}val')
        if val and val.isdigit():
            val_int = int(val)
            if val_int > 100:
                new_val = max(100, int(val_int * factor))
                trHeight.set(f'{{{ns}}}val', str(new_val))
                modified += 1
    
    print(f"  Shrunk {modified} Elemente")
    
    new_xml = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone='yes')
    
    tmp_path = docx_path.with_suffix('.tmp.docx')
    with zipfile.ZipFile(docx_path, 'r') as z_in:
        with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as z_out:
            for item in z_in.namelist():
                if item == 'word/document.xml':
                    z_out.writestr(item, new_xml)
                else:
                    z_out.writestr(item, z_in.read(item))
    
    import os
    os.replace(tmp_path, docx_path)

def update_checkboxes(docx_path, abteilung):
    """Setzt Checkboxen direkt im XML."""
    abt_lower = abteilung.lower()
    
    with zipfile.ZipFile(docx_path, 'r') as z:
        xml_content = z.read('word/document.xml')
    
    root = etree.fromstring(xml_content)
    
    # Checkboxen zurücksetzen
    for t in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
        if t.text in ['☐', '☒', '[X]', '[ ]']:
            t.text = '☐'
    
    # Richtige Checkbox aktivieren
    checkbox_index = 0
    for t in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
        if t.text == '☐':
            if abt_lower == 'betrieb' and checkbox_index == 0:
                t.text = '☒'
                break
            elif abt_lower == 'berufsschule' and checkbox_index == 1:
                t.text = '☒'
                break
            checkbox_index += 1
    
    new_xml = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone='yes')
    
    tmp_path = docx_path.with_suffix('.tmp.docx')
    with zipfile.ZipFile(docx_path, 'r') as z_in:
        with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as z_out:
            for item in z_in.namelist():
                if item == 'word/document.xml':
                    z_out.writestr(item, new_xml)
                else:
                    z_out.writestr(item, z_in.read(item))
    
    import os
    os.replace(tmp_path, docx_path)

def convert_to_pdf(docx_path, pdf_path):
    """Konvertiert DOCX zu PDF mit LibreOffice (100% Konsistenz)."""
    import os
    os.chdir(docx_path.parent)
    
    result = subprocess.run(
        ['libreoffice', '--headless', '--convert-to', 'pdf', str(docx_path), '--outdir', str(docx_path.parent)],
        capture_output=True, text=True, timeout=60
    )
    
    # LibreOffice benennt die Datei möglicherweise um
    expected_pdf = docx_path.with_suffix('.pdf')
    
    # Prüfe ob PDF existiert
    if expected_pdf.exists():
        return True
    
    # Alternativ: Suche nach generierter PDF
    for f in docx_path.parent.glob('*.pdf'):
        if f.stat().st_mtime > docx_path.stat().st_mtime:
            if f != pdf_path:
                f.rename(pdf_path)
            return True
    
    return False

def check_pages(pdf_path):
    """Gibt Seitenanzahl zurück."""
    result = subprocess.run(
        ['pdfinfo', str(pdf_path)],
        capture_output=True, text=True
    )
    for line in result.stdout.split('\n'):
        if line.startswith('Pages:'):
            return int(line.split(':')[1].strip())
    return -1

def main():
    parser = argparse.ArgumentParser(description='Berichtsheft DOCX Updater')
    parser.add_argument('--kw', type=int, required=True)
    parser.add_argument('--jahr', type=int, required=True)
    parser.add_argument('--nr', type=int, required=True)
    parser.add_argument('--abteilung', type=str, required=True, choices=['Betrieb', 'Berufsschule'])
    parser.add_argument('--montag', type=str, default='')
    parser.add_argument('--dienstag', type=str, default='')
    parser.add_argument('--mittwoch', type=str, default='')
    parser.add_argument('--donnerstag', type=str, default='')
    parser.add_argument('--freitag', type=str, default='')
    parser.add_argument('--thema', type=str, default='')
    parser.add_argument('--datum', type=str, default='')
    parser.add_argument('--ausbildungsjahr', type=int, default=3, help='Ausbildungsjahr (Standard: 3)')
    
    args = parser.parse_args()
    
    # Pfade
    kw_dir = Path.home() / 'Nextcloud' / 'Dokumente' / 'Ausbildung-Cancom' / 'Berichtsheft' / str(args.jahr) / f'KW{args.kw:02d}'
    kw_dir.mkdir(parents=True, exist_ok=True)
    docx_path = kw_dir / f'Berichtsheft Täglich KW {args.kw:02d}.docx'
    pdf_path = kw_dir / f'Berichtsheft Täglich KW {args.kw:02d}.pdf'
    
    if not docx_path.exists():
        print(f"FEHLER: DOCX nicht gefunden: {docx_path}")
        sys.exit(1)
    
    # Prüfe ob Platzhalter vorhanden sind
    placeholders_exist = has_placeholders(docx_path)
    
    if placeholders_exist:
        print(f"Öffne: {docx_path}")
        
        # Signaturdatum
        if args.datum:
            parts = args.datum.split(' - ')
            if len(parts) == 2:
                date_str = parts[1].strip()
                day_month = date_str.split('.')[0] + '.' + date_str.split('.')[1]
            else:
                day_month = "08.03"
        else:
            day_month = "08.03"
        
        signature_date = f"{day_month}.{args.jahr}"
        
        # Replacements
        replacements = {
            '{Day-of-signature}.{JAHRESJAHR}': signature_date,
            '{day-of.signature}.{JAHRESJAHR}': signature_date,
            '{Abteilung}': args.abteilung,
            '{Montag}': args.montag,
            '{Dienstag}': args.dienstag,
            '{Mittwoch}': args.mittwoch,
            '{Donnerstag}': args.donnerstag,
            '{Freitag}': args.freitag,
            '{week_topic}': args.thema,
            '{DATUMSZAHLUNG}': args.datum,
            '{WOCHENNUMMER}': f"KW {args.kw:02d}",
            '{AUSBILDUNGSJAHR}': str(args.ausbildungsjahr),
            '{Ausbildungsjahr}': str(args.ausbildungsjahr),
        }
        
        doc = Document(str(docx_path))
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_cell(cell, replacements)
        
        for para in doc.paragraphs:
            replace_in_paragraph(para, replacements)
        
        for para in doc.paragraphs:
            if 'Ausbildungsnachweis Nr.' in para.text:
                for run in para.runs:
                    if 'Nr.' in run.text:
                        run.text = re.sub(r'Nr\.\s*\d+', f"Nr. {args.nr}", run.text)
                break
        
        doc.save(str(docx_path))
        print(f"✓ Platzhalter ersetzt")
    
    # Checkboxen aktualisieren
    update_checkboxes(docx_path, args.abteilung)
    print(f"✓ Checkbox '{args.abteilung}' gesetzt")
    
    # PDF konvertieren und prüfen
    if not convert_to_pdf(docx_path, pdf_path):
        print("⚠️  PDF-Konvertierung fehlgeschlagen")
        sys.exit(0)
    
    pages = check_pages(pdf_path)
    
    if pages == 1:
        print(f"✓ Perfekt: 1 Seite")
        return
    
    if pages > 1:
        print(f"\n⚠️  MEHR ALS EINE SEITE ({pages} Seiten) - Verkleinere Zeilen...")
        
        factor = 0.95
        shrink_count = 0
        max_shrinks = 10
        
        while shrink_count < max_shrinks:
            shrink_count += 1
            
            shrink_xml(docx_path, factor=factor)
            update_checkboxes(docx_path, args.abteilung)
            
            if convert_to_pdf(docx_path, pdf_path):
                pages2 = check_pages(pdf_path)
                
                if pages2 == 1:
                    pct = int((1-factor)*100)
                    print(f"✓ Perfekt: 1 Seite (nach {shrink_count}. Anpassung, {pct}% Verkleinert)")
                    return
                
                if shrink_count >= 10:
                    print(f"⚠️  10x verkleinert, aber immernoch {pages2} Seiten!")
                    print("   → Text zu lang, bitte kürzen!")
                    return
                
                shrinkage_pct = int((1-factor)*100)
                if shrinkage_pct > 10:
                    print(f"⚠️  Schrumpfung >10% ({shrinkage_pct}%), Text zu lang!")
                else:
                    print(f"⚠️  Immernoch {pages2} Seiten - Schrumpfung {shrink_count}: {shrinkage_pct}%")
                
                factor = max(0.80, factor - 0.03)
            else:
                print("PDF-Konvertierung fehlgeschlagen")
                return

if __name__ == '__main__':
    main()
